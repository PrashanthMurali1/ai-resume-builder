import io, json, os, logging
from typing import List, Optional

import httpx
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from .prompts import TAILOR_PROMPT, KEYWORDS_PROMPT, COMPANY_PROMPT
from .utils import slugify, missing_keywords

load_dotenv()

logger = logging.getLogger("resume_tailor")
logging.basicConfig(level=os.getenv("LOG_LEVEL", "INFO"))

OLLAMA_URL   = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:1b")
OLLAMA_TIMEOUT = float(os.getenv("OLLAMA_TIMEOUT", "180"))
OLLAMA_DEBUG = os.getenv("OLLAMA_DEBUG", "0") not in ("0", "false", "False", "")

app = FastAPI(title="Resume Tailor Backend (Local Ollama)")

logger.info("Startup config: OLLAMA_MODEL=%s OLLAMA_URL=%s PID=%s", OLLAMA_MODEL, OLLAMA_URL, os.getpid())

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:19006", "http://localhost:8081", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Core helpers ----------
async def _ensure_model_available() -> None:
    """Check that the configured model appears in /api/tags list.
    Raises HTTPException(502) with a helpful message if not present.
    """
    model = OLLAMA_MODEL
    tags_url = f"{OLLAMA_URL}/api/tags"
    try:
        async with httpx.AsyncClient(timeout=OLLAMA_TIMEOUT) as client:
            r = await client.get(tags_url)
    except Exception as e:
        raise HTTPException(502, f"Cannot reach Ollama at {tags_url}: {e}")
    if r.status_code != 200:
        raise HTTPException(502, f"Ollama /api/tags returned {r.status_code}: {r.text[:200]}")
    try:
        data = r.json()
    except Exception:
        raise HTTPException(502, f"Invalid JSON from /api/tags: {r.text[:200]}")
    models = [m.get('name') for m in data.get('models', []) if isinstance(m, dict)]
    if model not in models:
        raise HTTPException(
            502,
            f"Model '{model}' not found in Ollama. Available: {', '.join(models) or 'NONE'}. "
            f"Pull it with: ollama pull {model.split(':')[0]}"
        )

async def ollama_generate(prompt: str) -> str:
    """Send a prompt to the local Ollama server and return its response.

    Converts any network / parsing issues into HTTP 502 so FastAPI does not emit a 500.
    """
    url = f"{OLLAMA_URL}/api/generate"
    payload = {"model": OLLAMA_MODEL, "prompt": prompt, "stream": False}
    # Proactively ensure model is present (cheap list call once if missing cached)
    # Only check again if we previously failed to call generate.
    try:
        async with httpx.AsyncClient(timeout=OLLAMA_TIMEOUT) as client:
            logger.info("Ollama generate model=%s prompt_len=%d", OLLAMA_MODEL, len(prompt))
            r = await client.post(url, json=payload)
    except Exception as e:
        logger.error("Ollama network failure: %s", e)
        # Double‑check tags to give a more helpful message (model missing vs network)
        await _ensure_model_available()
        raise HTTPException(502, f"Ollama request failed: {e}")

    if r.status_code != 200:
        logger.warning("Ollama non-200 (%s): %.300s", r.status_code, r.text)
        # Try to distinguish missing model vs generic error
        snippet = r.text[:300]
        if 'model' in snippet.lower() and 'not found' in snippet.lower():
            await _ensure_model_available()
        raise HTTPException(502, detail={"stage":"upstream-status","status":r.status_code,"snippet":snippet})

    # Parse JSON safely
    try:
        data = r.json()
    except Exception as e:
        logger.error("Ollama JSON parse error: %s; body=%.300s", e, r.text)
        raise HTTPException(502, detail={"stage":"json-parse","error":str(e),"body_head":r.text[:300]})

    if 'error' in data and data['error']:
        err = str(data['error'])
        logger.error("Ollama error field: %s", err)
        raise HTTPException(502, detail={"stage":"error-field","error": err[:300]})
    resp = data.get("response", "").strip()
    if not resp:
        logger.error("Empty response from Ollama; raw keys=%s", list(data.keys()))
        # If debug is on, surface entire upstream JSON
        detail = {"stage":"empty-response","keys":list(data.keys()),"raw": data if OLLAMA_DEBUG else None}
        raise HTTPException(502, detail=detail)
    return resp

def parse_docx_bytes(b: bytes) -> str:
    bio = io.BytesIO(b)
    doc = Document(bio)
    return "\n".join(p.text for p in doc.paragraphs)

def parse_pdf_bytes(b: bytes) -> str:
    bio = io.BytesIO(b)
    reader = PdfReader(bio)
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunks)

# ---------- Schemas ----------
class TailorReq(BaseModel):
    resume_text: str
    jd_text: str

class KeywordsReq(BaseModel):
    resume_text: str
    jd_text: str

class ExportReq(BaseModel):
    text: str
    label: Optional[str] = None      # e.g., "microsoft" -> resume_microsoft.docx/pdf
    fmt: str = "docx"                # "docx" or "pdf"

class ParseLocalReq(BaseModel):
    path: str

# ---------- Endpoints ----------
@app.get("/health")
async def health():
    info = {"model": OLLAMA_MODEL, "ollama_ok": False, "model_present": False}
    try:
        async with httpx.AsyncClient(timeout=4) as client:
            resp = await client.get(f"{OLLAMA_URL}/api/tags")
        info["ollama_ok"] = resp.status_code == 200
        if info["ollama_ok"]:
            try:
                tags = resp.json()
                models = [m.get('name') for m in tags.get('models', []) if isinstance(m, dict)]
                info["model_present"] = OLLAMA_MODEL in models
            except Exception:
                pass
    except Exception as e:
        info["error"] = str(e)
    return info

@app.post("/parse")
async def parse(file: UploadFile = File(...)):
    data = await file.read()
    name = (file.filename or "").lower()
    try:
        if name.endswith(".docx"):
            text = parse_docx_bytes(data)
        elif name.endswith(".pdf"):
            text = parse_pdf_bytes(data)
        else:
            text = data.decode("utf-8", errors="ignore")
        return {"text": text}
    except Exception as e:
        raise HTTPException(400, f"Failed to parse file: {e}")

@app.post("/tailor")
async def tailor(req: TailorReq):
    if not req.resume_text.strip():
        raise HTTPException(400, "resume_text is empty")
    if not req.jd_text.strip():
        raise HTTPException(400, "jd_text is empty")
    try:
        prompt = TAILOR_PROMPT.format(resume=req.resume_text, jd=req.jd_text)
    except KeyError as e:
        # In case braces in prompt cause format issues
        raise HTTPException(500, f"Prompt template error: missing key {e}")
    logger.info("/tailor resume_len=%d jd_len=%d", len(req.resume_text), len(req.jd_text))
    text = await ollama_generate(prompt)
    return {"tailored": text, "model": OLLAMA_MODEL}

@app.post("/debug/generate")
async def debug_generate(body: dict):
    """Forward a raw prompt to Ollama and return upstream payload (for troubleshooting 502s)."""
    prompt = body.get("prompt", "")
    model = body.get("model", OLLAMA_MODEL)
    if not prompt.strip():
        raise HTTPException(400, "prompt required")
    url = f"{OLLAMA_URL}/api/generate"
    payload = {"model": model, "prompt": prompt, "stream": False}
    try:
        async with httpx.AsyncClient(timeout=OLLAMA_TIMEOUT) as client:
            r = await client.post(url, json=payload)
    except Exception as e:
        return JSONResponse(status_code=502, content={"stage":"network","error":str(e)})
    head = r.text[:800]
    try:
        parsed = r.json()
    except Exception:
        parsed = None
    return {
        "status": r.status_code,
        "head": head,
        "json": parsed,
        "model": model,
        "prompt_len": len(prompt)
    }

@app.post("/keywords")
async def keywords(req: KeywordsReq):
    kw_json = await ollama_generate(KEYWORDS_PROMPT.format(jd=req.jd_text))
    # try to parse JSON array; fall back to naive split
    try:
        kws: List[str] = json.loads(kw_json)
        if not isinstance(kws, list):
            kws = []
    except Exception:
        kws = [k.strip("-• \n") for k in kw_json.split("\n") if k.strip()]
    missing = missing_keywords(req.resume_text, kws)
    return {"keywords": kws, "missing": missing}

@app.post("/parse-local")
async def parse_local(req: ParseLocalReq):
    try:
        path = req.path
        # Accept paths with optional file:// prefix
        if path.startswith("file://"):
            path = path[len("file://"):]
        with open(path, "rb") as f:
            data = f.read()
        name = path.lower()
        if name.endswith(".docx"):
            text = parse_docx_bytes(data)
        elif name.endswith(".pdf"):
            text = parse_pdf_bytes(data)
        else:
            text = data.decode("utf-8", errors="ignore")
        return {"text": text}
    except Exception as e:
        raise HTTPException(400, f"Failed to parse file: {e}")

@app.post("/infer-company")
async def infer_company(jd_text: str = Form(...)):
    name = (await ollama_generate(COMPANY_PROMPT.format(jd=jd_text))).strip()
    return {"company": name}

@app.post("/export")
async def export(res: ExportReq):
    base = f"resume_{slugify(res.label or 'draft')}"
    if res.fmt == "docx":
        doc = Document()
        for para in res.text.split("\n"):
            doc.add_paragraph(para)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return StreamingResponse(
            bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base}.docx"'}
        )
    elif res.fmt == "pdf":
        bio = io.BytesIO()
        c = canvas.Canvas(bio, pagesize=LETTER)
        width, height = LETTER
        left, top, leading = 72, height - 72, 14
        x, y = left, top
        for line in res.text.split("\n"):
            for chunk in [line[i:i+100] for i in range(0, len(line), 100)]:
                if y < 72:
                    c.showPage()
                    y = top
                c.drawString(x, y, chunk)
                y -= leading
        c.save()
        bio.seek(0)
        return StreamingResponse(
            bio, media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{base}.pdf"'}
        )
    else:
        raise HTTPException(400, "fmt must be 'docx' or 'pdf'")
